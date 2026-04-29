"""Chargement de documents: PDF, DOCX, TXT, MD, HTML, URL."""
import logging
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass
import urllib.request
import urllib.parse

logger = logging.getLogger(__name__)


@dataclass
class Document:
    content: str
    metadata: Dict[str, Any]

def _load_excel(path: str) -> str:
    """Charge un fichier Excel — fallback texte pour les formats simples."""
    docs = load_excel_as_documents(path)
    return "\n\n".join(d.content for d in docs)


def load_excel_as_documents(path: str) -> List[Document]:
    """Charge un fichier Excel et retourne UN Document par ligne.
    Chaque entrée est un texte autonome avec les noms de colonnes,
    ce qui permet au retriever de les associer sémantiquement
    et au chunker de ne pas les couper arbitrairement."""
    import openpyxl
    p = Path(path)
    wb = openpyxl.load_workbook(path, data_only=True)
    docs = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        # Lire les headers
        headers = []
        for cell in ws[1]:
            headers.append(str(cell.value).strip() if cell.value else "")

        for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 2):
            values = [str(v).strip() if v is not None else "" for v in row]
            if not any(values):
                continue
            # Construire un texte structuré "Colonne: Valeur" par ligne
            entry_parts = []
            for header, value in zip(headers, values):
                if header and value:
                    entry_parts.append(f"{header}: {value}")
            if not entry_parts:
                continue
            # Préfixer avec le nom du fichier source (sans extension)
            # pour aider le retriever et le LLM à distinguer les sources
            source_label = p.stem.upper()  # DIRECTION, DEPARTEMENT, SERVICE, POSTE
            content = f"[{source_label}] " + " | ".join(entry_parts)
            docs.append(Document(
                content=content,
                metadata={
                    "source": str(p),
                    "filename": p.name,
                    "extension": p.suffix.lower(),
                    "sheet": sheet_name,
                    "row": row_idx,
                    "size_bytes": p.stat().st_size,
                },
            ))
    return docs

def _load_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())


def _load_docx(path: str) -> str:
    from docx import Document as DocxDoc
    doc = DocxDoc(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _load_html(path: str) -> str:
    from bs4 import BeautifulSoup
    with open(path, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _load_text(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".doc": _load_docx,
    ".html": _load_html,
    ".htm": _load_html,
    ".txt": _load_text,
    ".md": _load_text,
    ".markdown": _load_text,
    ".rst": _load_text,
    ".csv": _load_text,
    ".json": _load_text,
      ".xlsx":     _load_excel,   # ← ajouté
    ".xls":      _load_excel,   # ← ajouté
}


def load_document(path: str) -> Document:
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in LOADERS:
        raise ValueError(f"Format non supporté: {ext}")

    content = LOADERS[ext](str(p))
    return Document(
        content=content,
        metadata={
            "source": str(p),
            "filename": p.name,
            "extension": ext,
            "size_bytes": p.stat().st_size,
        },
    )


def scrape_url(url: str, timeout: int = 15) -> Document:
    """Scrape une URL HTTP/HTTPS et retourne un Document."""
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError(f"URL invalide (http/https requis): {url}")

    req = urllib.request.Request(url, headers={"User-Agent": "RAGBot/1.0"})
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        raw_html = resp.read()

    from bs4 import BeautifulSoup
    soup = BeautifulSoup(raw_html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    title = soup.title.string.strip() if soup.title and soup.title.string else parsed.netloc
    content = soup.get_text(separator="\n", strip=True)

    # Nom de fichier safe pour les IDs ChromaDB
    safe_name = urllib.parse.quote(url, safe="").replace("%", "_")[:80] + ".url"

    return Document(
        content=content,
        metadata={
            "source": url,
            "filename": safe_name,
            "extension": ".url",
            "title": title,
        },
    )


def load_directory(directory: str) -> List[Document]:
    docs = []
    excel_exts = {".xlsx", ".xls"}
    for file in sorted(Path(directory).rglob("*")):
        if not file.is_file():
            continue
        ext = file.suffix.lower()
        if ext not in LOADERS:
            continue
        try:
            if ext in excel_exts:
                # Excel : un Document par ligne pour un meilleur retrieval
                excel_docs = load_excel_as_documents(str(file))
                docs.extend(excel_docs)
                logger.info("  Chargé: %s (%d entrées Excel)", file.name, len(excel_docs))
            else:
                doc = load_document(str(file))
                if doc.content.strip():
                    docs.append(doc)
                    logger.info("  Chargé: %s (%d chars)", file.name, len(doc.content))
        except Exception as e:
            logger.warning("  Ignoré %s: %s", file.name, e)
    return docs
